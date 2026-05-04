"""
Service layer for resume-related operations.
"""

import base64
import logging
import os
import time
from io import BytesIO
from typing import Any

import fitz
from django.core.files.uploadedfile import InMemoryUploadedFile
from docx import Document
from rest_framework.response import Response

from apps.resumes.exceptions import OpenAIServiceError, PDFGenerationError, ResumeProcessingError
from apps.resumes.services.ai.openai.client import OpenAIService
from apps.resumes.services.ai.openai.data_structures import get_default_resume_structure
from apps.resumes.services.resume.pdf import PDFService

logger = logging.getLogger(__name__)


class ResumeService:
    """Service class for resume operations."""

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

    def __init__(self):
        self.openai_service = OpenAIService()
        self.pdf_service = PDFService()

    def _get_file_extension(self, filename: str) -> str:
        """Get the file extension in lowercase."""
        return os.path.splitext(filename)[1].lower()

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        doc = Document(BytesIO(file_content))
        text_parts = []

        # Extract headers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)

        # Extract main document paragraphs (includes all text, headlines, etc.)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        # Extract footers
        for section in doc.sections:
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)

        return "\n".join(text_parts)

    def _extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from file based on its extension.

        Args:
            file_content: The file content as bytes
            filename: The name of the file (used to determine type)

        Returns:
            Extracted text content

        Raises:
            ResumeProcessingError: If file type is unsupported or extraction fails
        """
        ext = self._get_file_extension(filename)

        # Validate extension first
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ResumeProcessingError(
                detail=f"Unsupported file type: {ext}. Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        try:
            if ext == ".pdf":
                return self._extract_text_from_pdf(file_content)
            elif ext == ".docx":
                return self._extract_text_from_docx(file_content)
        except Exception as e:
            logger.error(f"Error extracting text from {ext}: {e!s}")
            raise ResumeProcessingError(detail=f"Failed to extract text from {ext.upper()}: {e!s}")

    def extract_resume_info(
        self, uploaded_file: InMemoryUploadedFile, language: str = "en"
    ) -> dict[str, Any]:
        """
        Extract resume information from an uploaded file.

        Args:
            uploaded_file: The uploaded file to process (PDF or DOCX)
            language: Language code for the output resume (default: 'en')

        Returns:
            Dict containing the extracted resume information

        Raises:
            OpenAIServiceError: If there's an error processing with the AI service
            ResumeProcessingError: If there's an error processing the file
        """
        start_time = time.time()
        logger.info(f"Starting resume info extraction for file: {uploaded_file.name}")

        try:
            # Read file content
            read_start = time.time()
            file_content = uploaded_file.read()
            read_time = time.time() - read_start
            logger.info(
                f"File read completed in {read_time:.2f} seconds, size: {len(file_content)} bytes"
            )

            # Extract text (validation happens inside this method)
            extract_start = time.time()
            text = self._extract_text_from_file(file_content, uploaded_file.name)
            extract_time = time.time() - extract_start
            logger.info(
                f"Text extraction completed in {extract_time:.2f} seconds, extracted {len(text)} characters"
            )

            # Call AI service
            ai_start = time.time()
            logger.info("Calling OpenAI service to extract structured data...")
            resume_data = self.openai_service.extract_resume_info(
                text, language=language, skip_post_processing=False
            )
            ai_time = time.time() - ai_start
            logger.info(f"OpenAI service completed in {ai_time:.2f} seconds")

            total_time = time.time() - start_time
            logger.info(
                f"Total resume extraction completed in {total_time:.2f} seconds "
                f"(read: {read_time:.2f}s, extract: {extract_time:.2f}s, ai: {ai_time:.2f}s)"
            )

            return resume_data

        except ResumeProcessingError:
            raise
        except Exception as e:
            total_time = time.time() - start_time
            logger.error(f"Error in resume extraction after {total_time:.2f} seconds: {e!s}")
            raise OpenAIServiceError(detail=f"Error extracting resume info: {e!s}")

    def generate_from_qa(
        self,
        qa_pairs: list[dict[str, str]],
        user_info: dict[str, str],
        language: str = "en",
        position: str | None = None,
    ) -> dict[str, Any]:
        """
        Generate resume data from Q&A pairs.

        Args:
            qa_pairs: List of question-answer pairs
            user_info: User information for the resume
            language: Language code for the output resume (default: 'en')
            position: Optional job position for the resume

        Returns:
            Dict containing the generated resume data

        Raises:
            OpenAIServiceError: If there's an error generating the resume
        """
        try:
            return self.openai_service.generate_resume_from_qa(
                qa_pairs, user_info, language, position
            )
        except Exception as e:
            raise OpenAIServiceError(detail=str(e))

    def generate_pdf(
        self,
        resume_data: dict[str, Any],
        profile_image: str | None = None,
    ) -> BytesIO:
        """
        Generate a PDF from resume data.

        Args:
            resume_data: Resume data to generate PDF from
            profile_image: Optional profile image (base64 or raw bytes)
            language: Language code for the PDF (default: 'en')

        Returns:
            BytesIO object containing the PDF

        Raises:
            PDFGenerationError: If there's an error generating the PDF
        """
        try:
            # Ensure resume_data has all expected keys (title, certifications, etc.)
            merged = get_default_resume_structure()
            if isinstance(resume_data, dict):
                merged.update(resume_data)

            if profile_image:
                if ";base64," in profile_image:
                    _, imgstr = profile_image.split(";base64,")
                    merged["profile_image"] = base64.b64decode(imgstr)
                else:
                    merged["profile_image"] = profile_image

            return self.pdf_service.generate_resume_pdf(merged)
        except Exception as e:
            raise PDFGenerationError(detail=str(e))

    @staticmethod
    def create_pdf_response(pdf_file: BytesIO, filename: str = "resume.pdf") -> Response:
        """
        Create a FileResponse from a PDF file.

        Args:
            pdf_file: BytesIO object containing the PDF
            filename: Filename for the download

        Returns:
            FileResponse with proper headers for PDF download
        """
        from django.http import FileResponse

        response = FileResponse(
            pdf_file, as_attachment=True, filename=filename, content_type="application/pdf"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response
